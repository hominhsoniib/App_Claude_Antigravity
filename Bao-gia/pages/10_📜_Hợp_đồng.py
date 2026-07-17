import streamlit as st
import pandas as pd
from database.db import get_session
from models.models import Customer, QuotationHeader, QuoteStatus
from auth.session import require_login
import importlib
from services import contract_ai_service as contract_ai
importlib.reload(contract_ai)
from io import BytesIO

st.set_page_config(page_title="Trợ lý Hợp đồng - QUOTEFLOW OS", page_icon="📜", layout="wide")

current_user = require_login()

st.title("📜 Trợ lý Hợp đồng AI")

# Xác định AI engine đang chạy
import os
from dotenv import load_dotenv
load_dotenv(override=True)
claude_key = os.getenv("CLAUDE_API_KEY", "").strip()
gemini_key = os.getenv("GEMINI_API_KEY", "").strip()
if claude_key and not claude_key.startswith("your-") and claude_key != "":
    ai_label = f"Claude ({os.getenv('CLAUDE_MODEL', 'claude-sonnet-5')})"
elif gemini_key and not gemini_key.startswith("your-") and gemini_key != "":
    ai_label = f"Gemini ({os.getenv('GEMINI_MODEL', 'gemini-1.5-flash')})"
else:
    ai_label = "Mô phỏng (Mock Engine)"

st.caption(f"Tư vấn soạn thảo hợp đồng chuyên nghiệp và rà soát rủi ro pháp lý sử dụng trí tuệ nhân tạo {ai_label}.")

# Kết nối CSDL để lấy danh sách Khách hàng & Báo giá
db = get_session()
customers = db.query(Customer).order_by(Customer.name).all()
# Lấy các báo giá
quotes = db.query(QuotationHeader).order_by(QuotationHeader.quote_no.desc()).all()


# Hàm tạo file DOCX từ văn bản hợp đồng để tải về
def generate_docx_download(contract_text, title):
    import docx
    doc = docx.Document()
    
    # Định dạng lề trang chuẩn văn bản hành chính (1 inch = 2.54 cm)
    for section in doc.sections:
        section.top_margin = docx.shared.Inches(1)
        section.bottom_margin = docx.shared.Inches(1)
        section.left_margin = docx.shared.Inches(1)
        section.right_margin = docx.shared.Inches(1)
        
    doc.add_heading(title, level=0)
    
    # Duyệt từng dòng để thêm vào paragraph
    lines = contract_text.split("\n")
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
        
        # Nhận diện tiêu đề đơn giản để format trong Word
        if line_clean.startswith("# "):
            doc.add_heading(line_clean.replace("# ", ""), level=1)
        elif line_clean.startswith("## "):
            doc.add_heading(line_clean.replace("## ", ""), level=2)
        elif line_clean.startswith("### "):
            doc.add_heading(line_clean.replace("### ", ""), level=3)
        elif line_clean.startswith("#### "):
            doc.add_heading(line_clean.replace("#### ", ""), level=4)
        else:
            doc.add_paragraph(line_clean)
            
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

tab1, tab2, tab3 = st.tabs(["✍️ Tư vấn & Soạn thảo Hợp đồng", "🔍 Rà soát & Sửa đổi Hợp đồng", "💬 Trò chuyện & Hỏi đáp"])

# ==========================================
# TAB 1: SOẠN THẢO HỢP ĐỒNG
# ==========================================
with tab1:
    st.subheader("Cung cấp thông tin hợp đồng")
    
    # Khởi tạo session state cho các trường nhập liệu nếu chưa có
    if "df_type" not in st.session_state:
        st.session_state["df_type"] = "Hợp đồng mua bán hàng hóa"
    if "df_seller" not in st.session_state:
        st.session_state["df_seller"] = "CÔNG TY CỔ PHẦN NEXUS CRM\nĐịa chỉ: Tòa nhà Nexus, 180 Nguyễn Thị Minh Khai, Quận 3, TP. Hồ Chí Minh\nMã số thuế: 0312345678\nĐại diện: Nguyễn Văn A - Chức vụ: Giám đốc"
    if "df_buyer" not in st.session_state:
        st.session_state["df_buyer"] = ""
    if "df_value" not in st.session_state:
        st.session_state["df_value"] = ""
    if "df_core" not in st.session_state:
        st.session_state["df_core"] = ""
    if "df_special" not in st.session_state:
        st.session_state["df_special"] = ""
    if "df_pay_method" not in st.session_state:
        st.session_state["df_pay_method"] = "Chuyển khoản qua tài khoản ngân hàng"
    if "df_pay_schedule" not in st.session_state:
        st.session_state["df_pay_schedule"] = "Bên B thanh toán tạm ứng 30% ngay sau khi ký hợp đồng; thanh toán 70% còn lại trong vòng 07 ngày làm việc kể từ ngày nghiệm thu bàn giao"
    if "df_lead_time" not in st.session_state:
        st.session_state["df_lead_time"] = "Trong vòng 10 ngày kể từ ngày ký hợp đồng"
    if "df_delivery_location" not in st.session_state:
        st.session_state["df_delivery_location"] = "Tại kho bên Mua"
    if "df_warranty" not in st.session_state:
        st.session_state["df_warranty"] = "12 tháng kể từ ngày nghiệm thu bàn giao"
    if "df_penalty" not in st.session_state:
        st.session_state["df_penalty"] = "Phạt chậm giao hàng 0.5% ngày chậm trễ, tối đa không quá 8% giá trị phần giao chậm"

    # Bảng điều khiển liên kết dữ liệu hệ thống
    with st.expander("🔗 Nhập liệu nhanh từ CRM (Khách hàng hoặc Báo giá)", expanded=True):
        c1, c2 = st.columns(2)
        with c1:
            quote_options = ["-- Chọn Báo giá --"] + [
                f"{q.quote_no} - {q.customer.name if q.customer else 'Khách vãng lai'} ({q.grand_total:,.0f} đ) [{q.status.value}]" 
                for q in quotes
            ]
            selected_quote = st.selectbox("Chọn từ danh sách Báo giá có sẵn:", options=quote_options, key="sel_quote_opt")
            
            if st.button("Áp dụng thông tin Báo giá này", use_container_width=True, type="secondary"):
                if selected_quote != "-- Chọn Báo giá --":
                    q_no = selected_quote.split(" - ")[0]
                    # Tìm lại báo giá trong CSDL
                    q_obj = db.query(QuotationHeader).filter(QuotationHeader.quote_no == q_no).first()
                    if q_obj:
                            st.session_state["df_type"] = "Hợp đồng mua bán hàng hóa"
                            # Nạp thông tin người mua từ khách hàng của báo giá
                            if q_obj.customer:
                                c = q_obj.customer
                                st.session_state["df_buyer"] = f"ĐƠN VỊ: {c.name}\nĐịa chỉ: {c.address or '[Địa chỉ]'}\nMã số thuế: {c.tax_code or '[MST]'}\nĐại diện: [Người đại diện] - Chức vụ: [Chức vụ]"
                            else:
                                st.session_state["df_buyer"] = ""
                                
                            st.session_state["df_value"] = f"{q_obj.grand_total:,.0f}"
                            
                            # Tổng hợp điều khoản cốt lõi từ danh mục sản phẩm của báo giá
                            prod_lines = []
                            for detail in q_obj.details:
                                p_name = detail.product.name if detail.product else "Sản phẩm"
                                unit = detail.product.unit if (detail.product and detail.product.unit) else "cái"
                                prod_lines.append(
                                    f"- {p_name}: {detail.qty} {unit} x {detail.unit_price:,.0f} đ (Chiết khấu {detail.discount_pct}%) = {detail.line_total:,.0f} đ"
                                )
                            
                            core_text = "Danh mục hàng hóa giao dịch:\n" + "\n".join(prod_lines)
                            st.session_state["df_pay_schedule"] = q_obj.payment_terms if q_obj.payment_terms else ""
                            st.session_state["df_delivery_location"] = q_obj.delivery_terms if q_obj.delivery_terms else "Tại kho bên Mua"
                            st.session_state["df_lead_time"] = q_obj.lead_time if q_obj.lead_time else "Trong vòng 10 ngày kể từ ngày ký hợp đồng"
                            
                            st.session_state["df_core"] = core_text
                            st.toast(f"✅ Đã tải thông tin từ Báo giá {q_no}!", icon="🎉")
                            db.close()
                            st.rerun()
                else:
                    st.warning("Vui lòng chọn một báo giá cụ thể trước.")
                    
        with c2:
            cust_options = ["-- Chọn Khách hàng --"] + [f"{c.code} - {c.name}" for c in customers]
            selected_cust = st.selectbox("Hoặc chỉ chọn thông tin Khách hàng:", options=cust_options, key="sel_cust_opt")
            
            if st.button("Áp dụng thông tin Khách hàng này", use_container_width=True, type="secondary"):
                if selected_cust != "-- Chọn Khách hàng --":
                    c_code = selected_cust.split(" - ")[0]
                    c_obj = db.query(Customer).filter(Customer.code == c_code).first()
                    if c_obj:
                            st.session_state["df_buyer"] = f"ĐƠN VỊ: {c_obj.name}\nĐịa chỉ: {c_obj.address or '[Địa chỉ]'}\nMã số thuế: {c_obj.tax_code or '[MST]'}\nĐại diện: [Người đại diện] - Chức vụ: [Chức vụ]"
                            st.toast(f"✅ Đã tải thông tin Khách hàng {c_obj.name}!", icon="🎉")
                            db.close()
                            st.rerun()
                else:
                    st.warning("Vui lòng chọn một khách hàng cụ thể trước.")

    # Form nhập chi tiết hợp đồng để gửi AI
    with st.form("contract_draft_form"):
        col1, col2 = st.columns(2)
        with col1:
            contract_type = st.selectbox(
                "Loại hợp đồng*",
                ["Hợp đồng mua bán hàng hóa", "Hợp đồng dịch vụ thương mại", "Hợp đồng hợp tác kinh doanh", "Hợp đồng thuê tài sản", "Hợp đồng lao động", "Khác"],
                index=["Hợp đồng mua bán hàng hóa", "Hợp đồng dịch vụ thương mại", "Hợp đồng hợp tác kinh doanh", "Hợp đồng thuê tài sản", "Hợp đồng lao động", "Khác"].index(st.session_state["df_type"])
            )
            seller_info = st.text_area(
                "Bên A (Bên bán/Bên cung cấp dịch vụ)*",
                value=st.session_state["df_seller"],
                height=120,
                help="Thông tin chi tiết của Bên A bao gồm Tên công ty, MST, Địa chỉ, Đại diện..."
            )
            buyer_info = st.text_area(
                "Bên B (Bên mua/Bên sử dụng dịch vụ)*",
                value=st.session_state["df_buyer"],
                height=120,
                placeholder="Nhập thông tin Bên B hoặc chọn nhập liệu nhanh ở trên...",
                help="Thông tin chi tiết của Bên B..."
            )
        with col2:
            value_str = st.text_input(
                "Giá trị hợp đồng (VNĐ)*",
                value=st.session_state["df_value"],
                placeholder="Ví dụ: 150,000,000"
            )
            core_terms = st.text_area(
                "Điều khoản cốt lõi (Sản phẩm, Thanh toán, Giao hàng)*",
                value=st.session_state["df_core"],
                height=120,
                placeholder="Ví dụ:\n- Bàn giao 10 Laptop Dell trước ngày 20/07\n- Thanh toán 30% khi ký hợp đồng, 70% sau khi bàn giao...",
                help="Các nội dung chính cần quy định cụ thể trong hợp đồng"
            )
            special_terms = st.text_area(
                "Yêu cầu bổ sung đặc biệt",
                value=st.session_state["df_special"],
                height=70,
                placeholder="Ví dụ: Cần điều khoản phạt giao hàng trễ 1%/ngày; Điều khoản bảo mật NDA 2 năm...",
                help="Bất cứ yêu cầu đặc thù nào bạn muốn AI đưa vào dự thảo hợp đồng"
            )
            
        # Thêm các điều khoản thương mại chi tiết
        st.write("")
        st.markdown("##### 📝 Điều khoản chi tiết (Thanh toán, giao nhận, phạt...)")
        col_c1, col_c2 = st.columns(2)
        with col_c1:
            payment_method = st.selectbox(
                "Phương thức thanh toán",
                ["Chuyển khoản qua tài khoản ngân hàng", "Thanh toán bằng tiền mặt", "Chuyển khoản hoặc Tiền mặt", "Khác"],
                index=["Chuyển khoản qua tài khoản ngân hàng", "Thanh toán bằng tiền mặt", "Chuyển khoản hoặc Tiền mặt", "Khác"].index(st.session_state["df_pay_method"]) if st.session_state["df_pay_method"] in ["Chuyển khoản qua tài khoản ngân hàng", "Thanh toán bằng tiền mặt", "Chuyển khoản hoặc Tiền mặt", "Khác"] else 0
            )
            payment_schedule = st.text_input(
                "Tiến độ thanh toán",
                value=st.session_state["df_pay_schedule"],
                placeholder="Ví dụ: Tạm ứng 30% ngay sau khi ký; 70% sau khi nghiệm thu..."
            )
            lead_time = st.text_input(
                "Thời hạn thực hiện / giao hàng",
                value=st.session_state["df_lead_time"],
                placeholder="Ví dụ: Trong vòng 10 ngày kể từ ngày ký..."
            )
        with col_c2:
            delivery_location = st.text_input(
                "Địa điểm thực hiện / giao hàng",
                value=st.session_state["df_delivery_location"],
                placeholder="Ví dụ: Tại kho bên Mua..."
            )
            warranty_period = st.text_input(
                "Thời hạn bảo hành",
                value=st.session_state["df_warranty"],
                placeholder="Ví dụ: 12 tháng kể từ ngày nghiệm thu bàn giao..."
            )
            penalty_terms = st.text_input(
                "Điều khoản phạt vi phạm",
                value=st.session_state["df_penalty"],
                placeholder="Ví dụ: Phạt vi phạm giao hàng trễ 0.5%/ngày, tối đa 8%..."
            )
            
        submitted = st.form_submit_button("✍️ Tạo dự thảo & Tư vấn soạn thảo", use_container_width=True, type="primary")

    if submitted:
        if not seller_info.strip() or not buyer_info.strip() or not core_terms.strip():
            st.error("❌ Vui lòng điền đầy đủ các trường thông tin đánh dấu hoa thị (*).")
        else:
            with st.spinner("🔄 AI đang soạn thảo dự thảo hợp đồng và phân tích pháp lý..."):
                response = contract_ai.draft_contract(
                    contract_type, seller_info, buyer_info, value_str, core_terms, special_terms,
                    payment_method, payment_schedule, lead_time, delivery_location, warranty_period, penalty_terms
                )
            
            # Tách phần văn bản hợp đồng và phần tư vấn của AI
            contract_text = ""
            advice_text = ""
            if "--- AI_ADVICE ---" in response:
                parts = response.split("--- AI_ADVICE ---")
                contract_text = parts[0].strip()
                advice_text = parts[1].strip()
            else:
                contract_text = response.strip()
                advice_text = "Không tìm thấy tư vấn riêng từ AI. Vui lòng tham khảo nội dung hợp đồng soạn thảo phía trên."
                
            st.session_state["generated_contract_text"] = contract_text
            st.session_state["generated_advice_text"] = advice_text
            st.session_state["generated_contract_title"] = f"DỰ THẢO {contract_type.upper()}"
            st.session_state["generated_contract_type"] = contract_type

    # Hiển thị kết quả soạn thảo
    if "generated_contract_text" in st.session_state:
        st.divider()
        st.subheader("🎉 Kết quả Soạn thảo từ AI")
        st.info(f"DEBUG: Chiều dài nội dung hợp đồng: {len(st.session_state['generated_contract_text'])} ký tự.")
        if len(st.session_state['generated_contract_text']) > 0:
            with st.expander("Xem nhanh nội dung gốc (Debug):", expanded=True):
                st.code(st.session_state['generated_contract_text'][:1000])
        
        c1, c2 = st.columns([7, 3])
        with c1:
            st.markdown("### 📜 Dự thảo Hợp đồng")
            # Text area để người dùng có thể chỉnh sửa lại trực tiếp
            editable_contract = st.text_area(
                "Nội dung Hợp đồng (Có thể chỉnh sửa thủ công trước khi tải về):", 
                value=st.session_state["generated_contract_text"],
                height=500
            )
            
            # Nút tải xuống
            d_col1, d_col2 = st.columns(2)
            with d_col1:
                docx_data = generate_docx_download(editable_contract, st.session_state["generated_contract_title"])
                st.download_button(
                    label="📥 Tải xuống file Word (.docx)",
                    data=docx_data,
                    file_name=f"Draft_{st.session_state['generated_contract_type'].replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary"
                )
            with d_col2:
                st.download_button(
                    label="📥 Tải xuống file Text (.txt)",
                    data=editable_contract,
                    file_name=f"Draft_{st.session_state['generated_contract_type'].replace(' ', '_')}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
                
        with c2:
            st.markdown("### 💡 Phân tích & Tư vấn của AI")
            st.info(st.session_state["generated_advice_text"])

# ==========================================
# TAB 2: RÀ SOÁT HỢP ĐỒNG DRAFT
# ==========================================
with tab2:
    st.subheader("Rà soát hợp đồng nháp & Đề xuất chỉnh sửa")
    st.write("Tải lên file hợp đồng hiện tại hoặc dán nội dung vào ô bên dưới để AI rà soát lỗi pháp lý, rủi ro điều khoản bất lợi.")

    # File uploader hỗ trợ txt, docx, pdf
    uploaded_file = st.file_uploader("Chọn file hợp đồng nháp (Hỗ trợ .txt, .docx, .pdf):", type=["txt", "docx", "pdf"])
    
    parsed_text = ""
    if uploaded_file is not None:
        file_details = {"FileName": uploaded_file.name, "FileType": uploaded_file.type}
        
        with st.spinner("🔄 Đang trích xuất nội dung từ file tải lên..."):
            try:
                if uploaded_file.name.endswith(".txt"):
                    parsed_text = uploaded_file.read().decode("utf-8", errors="ignore")
                    st.success(f"✅ Đã đọc thành công file Text: {uploaded_file.name}")
                elif uploaded_file.name.endswith(".docx"):
                    import docx
                    doc = docx.Document(uploaded_file)
                    parsed_text = "\n".join([para.text for para in doc.paragraphs if para.text])
                    st.success(f"✅ Đã đọc thành công file Word (.docx): {uploaded_file.name}")
                elif uploaded_file.name.endswith(".pdf"):
                    import pypdf
                    reader = pypdf.PdfReader(uploaded_file)
                    pages = []
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            pages.append(text)
                    parsed_text = "\n".join(pages)
                    st.success(f"✅ Đã đọc thành công file PDF: {uploaded_file.name}")
            except Exception as e:
                st.error(f"❌ Có lỗi khi đọc file: {str(e)}")
                
    # Dán văn bản trực tiếp
    input_text = st.text_area(
        "Nội dung hợp đồng cần rà soát (Có thể tự chỉnh sửa hoặc dán thêm ở đây):",
        value=parsed_text,
        placeholder="Dán toàn bộ nội dung hợp đồng nháp vào đây...",
        height=300,
        key="review_contract_input"
    )

    with st.form("contract_review_form"):
        r_col1, r_col2 = st.columns(2)
        with r_col1:
            user_role = st.selectbox(
                "Rà soát dưới góc độ bảo vệ quyền lợi của ai?*",
                ["Bên mua / Bên sử dụng dịch vụ (Bảo vệ người trả tiền)", "Bên bán / Bên cung cấp dịch vụ (Bảo vệ người bán hàng)"]
            )
        with r_col2:
            key_concerns = st.text_input(
                "Các điểm bạn lo ngại nhất (nếu có):",
                placeholder="Ví dụ: Muốn kiểm tra điều khoản phạt chậm thanh toán và bảo hành..."
            )
            
        review_submitted = st.form_submit_button("🔍 Rà soát & Phát hiện rủi ro điều khoản", use_container_width=True, type="primary")

    if review_submitted:
        if not input_text.strip():
            st.error("❌ Vui lòng cung cấp nội dung hợp đồng (bằng cách tải file lên hoặc dán nội dung trực tiếp).")
        else:
            with st.spinner("🔄 AI đang rà soát hợp đồng nháp và chuẩn bị đề xuất điều chỉnh..."):
                role_label = "Bên mua" if "Bên mua" in user_role else "Bên bán"
                review_result = contract_ai.review_contract(input_text, role_label, key_concerns)
                st.session_state["generated_review_result"] = review_result

    if "generated_review_result" in st.session_state:
        st.divider()
        st.subheader("📝 Báo cáo Rà soát Pháp lý từ AI")
        st.markdown(st.session_state["generated_review_result"])

# ==========================================
# TAB 3: TRÒ CHUYỆN & HỎI ĐÁP
# ==========================================
with tab3:
    st.subheader("💬 Trò chuyện & Hỏi đáp Hợp đồng")
    st.caption("Trao đổi trực tiếp với AI để chỉnh sửa, làm rõ các điều khoản, hoặc hỏi bất kỳ câu hỏi nào về pháp lý hợp đồng.")
    
    # Xác định ngữ cảnh hợp đồng hiện tại để gửi kèm
    current_contract_context = ""
    source_label = ""
    
    # Kiểm tra xem có bản nháp từ Tab 1 hay Tab 2 không
    if "generated_contract_text" in st.session_state and st.session_state["generated_contract_text"].strip():
        current_contract_context = st.session_state["generated_contract_text"]
        source_label = "Dự thảo hợp đồng từ Tab 1"
    elif "review_contract_input" in st.session_state and st.session_state["review_contract_input"].strip():
        current_contract_context = st.session_state["review_contract_input"]
        source_label = "Nội dung hợp đồng từ Tab 2"
        
    if current_contract_context:
        st.info(f"ℹ️ **Ngữ cảnh hiện tại:** AI đang tự động tham chiếu đến **{source_label}** để hỗ trợ bạn thảo luận.")
    else:
        st.warning("💡 **Gợi ý:** Bạn chưa soạn thảo ở Tab 1 hoặc dán hợp đồng ở Tab 2. Bạn có thể trò chuyện hỏi đáp các câu hỏi chung, hoặc nạp hợp đồng trước để AI hỗ trợ sát sườn nhất.")
        
    # Gợi ý một số câu hỏi nhanh
    st.write("**Câu hỏi gợi ý nhanh:**")
    chat_sample_questions = [
        "Mức phạt vi phạm tối đa theo quy định là bao nhiêu?",
        "Tư vấn điều khoản thanh toán 3 đợt an toàn?",
        "Bổ dung điều khoản bảo mật thông tin NDA?",
        "Lưu ý pháp lý quan trọng khi làm hợp đồng dịch vụ?",
    ]
    
    cols = st.columns(len(chat_sample_questions))
    clicked_chat_q = None
    for i, q in enumerate(chat_sample_questions):
        if cols[i].button(q, use_container_width=True, key=f"chat_q_{i}"):
            clicked_chat_q = q
            
    # Khởi tạo lịch sử chat hợp đồng
    if "contract_chat_history" not in st.session_state:
        st.session_state["contract_chat_history"] = []
        
    chat_input_q = st.chat_input("Nhập câu hỏi hoặc yêu cầu chỉnh sửa hợp đồng của bạn...")
    final_chat_q = clicked_chat_q or chat_input_q
    
    if final_chat_q:
        with st.spinner("🔄 AI đang suy nghĩ câu trả lời..."):
            ans = contract_ai.chat_about_contract(final_chat_q, current_contract_context, st.session_state["contract_chat_history"])
        st.session_state["contract_chat_history"].append({"q": final_chat_q, "a": ans})
        db.close()
        st.rerun()
        
    # Hiển thị lịch sử chat
    if st.session_state["contract_chat_history"]:
        if st.button("🗑️ Xóa lịch sử chat", type="secondary"):
            st.session_state["contract_chat_history"] = []
            db.close()
            st.rerun()
            
        for item in reversed(st.session_state["contract_chat_history"]):
            with st.chat_message("user"):
                st.write(item["q"])
            with st.chat_message("assistant"):
                st.markdown(item["a"])

# Đóng session ở cuối trang
db.close()

