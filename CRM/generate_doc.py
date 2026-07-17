import os
import re
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    """Đặt màu nền cho ô trong bảng (Shading)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def create_guide_from_md():
    md_path = os.path.join(os.path.dirname(__file__), "Huong_dan_su_dung_NEXUS_CRM.md")
    dest_path = os.path.join(os.path.dirname(__file__), "Huong_dan_su_dung_NEXUS_CRM.docx")
    
    if not os.path.exists(md_path):
        print(f"Lỗi: Không tìm thấy file markdown hướng dẫn: {md_path}")
        return
        
    doc = docx.Document()
    
    # Kiểu chữ mặc định
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Thiết lập lề trang
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    color_primary = RGBColor(31, 56, 100)   # #1F3864 - Navy Đậm
    color_secondary = RGBColor(37, 99, 235) # #2563EB - Xanh Blue
    color_dark = RGBColor(15, 23, 42)       # #0F172A - Xám đen
    color_muted = RGBColor(71, 85, 105)     # #475569 - Ghi nhạt
    
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    in_code_block = False
    code_text = []
    
    in_table = False
    table_data = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        line_stripped = line.strip()
        
        # 1. Xử lý Code Block
        if line_stripped.startswith("```"):
            if in_code_block:
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                
                run = p.add_run("\n".join(code_text))
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)
                run.font.color.rgb = color_muted
                code_text = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_text.append(line.rstrip('\n'))
            i += 1
            continue
            
        # 2. Xử lý Bảng biểu (Table)
        if line_stripped.startswith("|"):
            in_table = True
            table_data.append(line_stripped)
            i += 1
            continue
        elif in_table:
            in_table = False
            valid_rows = []
            for r in table_data:
                if re.match(r'^\|\s*[:\-|\s]+$', r):
                    continue
                valid_rows.append(r)
                
            if valid_rows:
                parsed_rows = []
                for vr in valid_rows:
                    cols = [c.strip() for c in vr.split("|")[1:-1]]
                    parsed_rows.append(cols)
                    
                if parsed_rows:
                    num_cols = len(parsed_rows[0])
                    num_rows = len(parsed_rows)
                    
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'
                    
                    for r_idx, row_cells in enumerate(parsed_rows):
                        for c_idx, val in enumerate(row_cells):
                            # Đảm bảo index cột không bị tràn
                            if c_idx >= len(table.rows[r_idx].cells):
                                continue
                            cell = table.cell(r_idx, c_idx)
                            # Dọn dẹp ký tự markdown
                            clean_val = val.replace("**", "").replace("`", "")
                            
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(3)
                            p.paragraph_format.space_after = Pt(3)
                            
                            run = p.add_run(clean_val)
                            run.font.size = Pt(10)
                            
                            if r_idx == 0:
                                set_cell_background(cell, "1F3864")
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                            else:
                                if r_idx % 2 == 1:
                                    set_cell_background(cell, "F8FAFC")
                                    
            table_data = []
            if not line_stripped:
                i += 1
                continue
                
        # 3. Dòng trống
        if not line_stripped:
            i += 1
            continue
            
        # 4. Tiêu đề lớn Heading 1 (#)
        if line_stripped.startswith("# "):
            title_text = line_stripped[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            
            run = p.add_run(title_text)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = color_primary
            i += 1
            continue
            
        # 5. Tiêu đề mục Heading 2 (##)
        if line_stripped.startswith("## "):
            h_text = line_stripped[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(h_text)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = color_primary
            i += 1
            continue
            
        # 6. Tiêu đề con Heading 3 (###)
        if line_stripped.startswith("### "):
            h_text = line_stripped[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(h_text)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = color_secondary
            i += 1
            continue
            
        # 7. Khung trích dẫn Blockquote (>)
        if line_stripped.startswith(">"):
            text = line_stripped[1:].strip()
            # Loại bỏ các tag alert
            text = re.sub(r'^\[!(NOTE|IMPORTANT|WARNING|TIP|CAUTION)\]\s*', '', text)
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = color_muted
            i += 1
            continue
            
        # 8. List items (Đầu dòng hoặc thứ tự)
        is_bullet = False
        is_num = False
        clean_line = line_stripped
        
        if line_stripped.startswith("- ") or line_stripped.startswith("* "):
            is_bullet = True
            clean_line = line_stripped[2:]
        elif re.match(r'^\d+\.\s', line_stripped):
            is_num = True
            match = re.match(r'^(\d+)\.\s(.*)', line_stripped)
            num_val = match.group(1)
            clean_line = match.group(2)
            
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        
        if is_bullet:
            p.style = 'List Bullet'
            p.paragraph_format.left_indent = Inches(0.25)
        elif is_num:
            p.paragraph_format.left_indent = Inches(0.25)
            run_num = p.add_run(f"{num_val}. ")
            run_num.bold = True
            run_num.font.color.rgb = color_secondary
            
        # Phân tích cú pháp các chữ in đậm, in nghiêng lồng nhau
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', clean_line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                r = p.add_run(part[2:-2])
                r.bold = True
            elif part.startswith("*") and part.endswith("*"):
                r = p.add_run(part[1:-1])
                r.italic = True
            elif part.startswith("`") and part.endswith("`"):
                r = p.add_run(part[1:-1])
                r.font.name = 'Courier New'
                r.font.size = Pt(9.5)
                r.font.color.rgb = color_muted
            else:
                p.add_run(part)
                
        i += 1
        
    doc.save(dest_path)
    print(f"[SUCCESS] Word document guide generated successfully at: {dest_path}")

if __name__ == "__main__":
    create_guide_from_md()
