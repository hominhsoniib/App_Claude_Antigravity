import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_guide():
    doc = docx.Document()

    # Cấu hình Margins (Lề trang)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Định dạng font chữ mặc định
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Hàm tạo tiêu đề tùy chỉnh
    def add_custom_heading(text, level, color=RGBColor(31, 56, 100)):
        h = doc.add_heading(level=level)
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.color.rgb = color
        if level == 1:
            run.font.size = Pt(16)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)
        elif level == 2:
            run.font.size = Pt(13)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
        return h

    # Trang bìa / Header
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("TÀI LIỆU HƯỚNG DẪN CẤU HÌNH & KÍCH HOẠT BẢN QUYỀN")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(31, 56, 100)
    title.paragraph_format.space_after = Pt(2)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("Hệ thống quản lý Báo giá và Hợp đồng - NEXUS CRM\nCơ chế ký số mã hóa RSA 2048-bit phục vụ chuyển giao phần mềm")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 100, 100)
    subtitle.paragraph_format.space_after = Pt(24)

    # --- Phần 1 ---
    add_custom_heading("1. Giới thiệu tổng quan hệ thống bản quyền", 1)
    p1 = doc.add_paragraph()
    p1.add_run("Hệ thống bản quyền của NEXUS CRM sử dụng công nghệ ký số không đối xứng RSA (thuật toán RS256 - mã hóa 2048 bit). Giải pháp này đảm bảo tính an toàn tuyệt đối khi chuyển giao phần mềm giữa các doanh nghiệp/đối tác mà không sợ bị chỉnh sửa thông tin bản quyền hoặc sao chép trái phép sang các máy chủ khác.")
    p1.paragraph_format.space_after = Pt(8)

    p2 = doc.add_paragraph()
    p2.add_run("Tệp tin bản quyền ")
    p2.add_run("license.lic").bold = True
    p2.add_run(" thực chất là một chuỗi mã hóa dạng JWT được ký số bằng Private Key (Khóa bí mật) lưu trữ tại NEXUS LICENSE SERVER và được xác thực ở phía Client thông qua Public Key (Khóa công khai) được nhúng sẵn trong ứng dụng. Mọi sự thay đổi dù chỉ 1 ký tự trong file ")
    p2.add_run("license.lic").bold = True
    p2.add_run(" sẽ lập tức làm mất hiệu lực chữ ký số và khóa hệ thống.")
    p2.paragraph_format.space_after = Pt(12)

    # --- Phần 2 ---
    add_custom_heading("2. Lấy thông tin vân tay phần cứng (Server Fingerprint)", 1)
    p3 = doc.add_paragraph()
    p3.add_run("Mã vân tay phần cứng (Server Fingerprint) là một chuỗi băm duy nhất đại diện cho cấu hình phần cứng máy chủ chạy ứng dụng (dựa trên BIOS UUID và địa chỉ MAC). Việc này giúp khóa cứng ứng dụng chỉ được chạy trên đúng máy chủ đã đăng ký.")
    p3.paragraph_format.space_after = Pt(8)

    p4 = doc.add_paragraph()
    p4.add_run("Để lấy mã vân tay máy chủ, người dùng thực hiện theo 2 cách:")
    p4.paragraph_format.space_after = Pt(4)

    doc.add_paragraph("- Cách 1: Truy cập trực tiếp vào giao diện ứng dụng -> Mục Cài đặt. Tại phần Bản quyền hệ thống, hệ thống sẽ tự động quét phần cứng và hiển thị mã Server Fingerprint (Ví dụ: E63495DC03165BB7). Sao chép chuỗi mã này gửi cho Nhà phát triển.", style='List Bullet')
    doc.add_paragraph("- Cách 2: Chạy trực tiếp file kiểm tra độc lập trên terminal của máy chủ để xuất mã định danh.", style='List Bullet')

    # --- Phần 3 ---
    add_custom_heading("3. Các bước sinh tệp tin bản quyền (Dành cho nhà phát triển)", 1)
    p5 = doc.add_paragraph()
    p5.add_run("Khi cần chuyển giao phần mềm sang một Công ty đối tác mới, Nhà phát triển thực hiện sinh tệp tin bản quyền theo các bước sau:")
    p5.paragraph_format.space_after = Pt(6)

    p6 = doc.add_paragraph()
    p6.add_run("Bước 3.1: ").bold = True
    p6.add_run("Mở terminal tại thư mục gốc của dự án và chạy tập lệnh sinh bản quyền:")
    p6.paragraph_format.space_after = Pt(4)

    code_p = doc.add_paragraph()
    code_p.paragraph_format.left_indent = Inches(0.5)
    code_run = code_p.add_run("python generate_license.py")
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10.5)
    code_run.font.bold = True
    code_run.font.color.rgb = RGBColor(180, 0, 0)
    code_p.paragraph_format.space_after = Pt(6)

    p7 = doc.add_paragraph()
    p7.add_run("Bước 3.2: ").bold = True
    p7.add_run("Nhập tuần tự các thông tin được yêu cầu trên Terminal:")
    p7.paragraph_format.space_after = Pt(4)

    doc.add_paragraph("1. Tên Công ty sở hữu: Tên hiển thị của đối tác (ví dụ: Công ty TNHH Giải pháp Công nghệ Việt).", style='List Number')
    doc.add_paragraph("2. Mã số thuế (MST): Mã số thuế đối tác sử dụng để đăng ký doanh nghiệp.", style='List Number')
    doc.add_paragraph("3. Tên miền bắt buộc: Định dạng email được quyền đăng nhập (ví dụ: @company.vn). Tất cả tài khoản có đuôi email này sẽ được cấp quyền.", style='List Number')
    doc.add_paragraph("4. Email đại diện: Địa chỉ email của quản trị viên (ví dụ: admin@company.vn).", style='List Number')
    doc.add_paragraph("5. Mã vân tay máy chủ: Dán mã Server Fingerprint đã lấy ở mục 2. Nếu muốn cho chạy trên mọi máy chủ mà không giới hạn phần cứng, nhập dấu hoa thị (*).", style='List Number')
    doc.add_paragraph("6. Số ngày hiệu lực: Thời gian sử dụng bản quyền tính bằng ngày (ví dụ: 365).", style='List Number')

    p8 = doc.add_paragraph()
    p8.add_run("Sau khi hoàn tất, hệ thống sẽ tự động xuất ra file ")
    p8.add_run("license.lic").bold = True
    p8.add_run(" đặt tại thư mục gốc của dự án. Hãy gửi file này cho khách hàng.")
    p8.paragraph_format.space_after = Pt(12)

    # --- Phần 4 ---
    add_custom_heading("4. Các bước cài đặt / Tải lên tệp tin bản quyền", 1)
    p9 = doc.add_paragraph()
    p9.add_run("Sau khi nhận được file ")
    p9.add_run("license.lic").bold = True
    p9.add_run(" từ nhà phát triển, khách hàng thực hiện cài đặt theo 2 cách:")
    p9.paragraph_format.space_after = Pt(6)

    p10 = doc.add_paragraph()
    p10.add_run("Cách 1: Tải lên qua giao diện quản trị (Khuyên dùng)").bold = True
    p10.paragraph_format.space_after = Pt(4)
    doc.add_paragraph("1. Đăng nhập vào hệ thống bằng tài khoản Quản trị viên (Admin).", style='List Number')
    doc.add_paragraph("2. Chọn mục Cài đặt ở menu bên trái.", style='List Number')
    doc.add_paragraph("3. Tại phần Bản quyền hệ thống, nhấn nút Browse files và chọn file license.lic đã nhận.", style='List Number')
    doc.add_paragraph("4. Giao diện hiển thị thông tin Công ty mới và tự động kích hoạt các tính năng tương ứng.", style='List Number')

    p11 = doc.add_paragraph()
    p11.add_run("Cách 2: Chép trực tiếp vào thư mục mã nguồn").bold = True
    p11.paragraph_format.space_after = Pt(4)
    p12 = doc.add_paragraph()
    p12.add_run("Chép file ")
    p12.add_run("license.lic").bold = True
    p12.add_run(" trực tiếp vào thư mục gốc ")
    p12.add_run("NEXUS-CRM/").italic = True
    p12.add_run(" hoặc thư mục ứng dụng ")
    p12.add_run("NEXUS-CRM/Bao-gia/").italic = True
    p12.add_run(". Hệ thống sẽ tự động nhận diện và cập nhật bản quyền ngay lập tức mà không cần khởi động lại dịch vụ.")
    p12.paragraph_format.space_after = Pt(12)

    # --- Phần 5 ---
    add_custom_heading("5. Phân quyền và Danh sách các Module hỗ trợ khóa bản quyền", 1)
    p13 = doc.add_paragraph()
    p13.add_run("Hệ thống hỗ trợ khóa và phân quyền linh hoạt theo từng module được định nghĩa sẵn trong tệp tin bản quyền:")
    p13.paragraph_format.space_after = Pt(8)

    # Thêm bảng thông tin Module
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Shading Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ký hiệu Module'
    hdr_cells[1].text = 'Tên Trang ứng dụng'
    hdr_cells[2].text = 'Mô tả tính năng'

    # Định dạng Bold cho header của bảng
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True

    modules_data = [
        ("quotation", "3_📝_Báo_giá.py, 4_✅_Phê_duyệt.py", "Quản lý tạo báo giá, quy trình duyệt đa cấp và in PDF mẫu báo giá chuyên nghiệp."),
        ("contract", "10_📜_Hợp_đồng.py", "Soạn thảo, rà soát và tư vấn hợp đồng song ngữ Anh-Việt tích hợp AI trợ lý pháp lý."),
        ("copilot", "6_🤖_AI_Copilot.py", "Trợ lý ảo thông minh hỗ trợ phân tích dữ liệu, tra cứu thông tin khách hàng bằng ngôn ngữ tự nhiên."),
        ("inventory", "9_📦_Tồn_kho.py", "Quản lý số lượng xuất nhập tồn sản phẩm, đồng bộ hóa danh mục kho vật tư xây dựng.")
    ]

    for i, (code, pages, desc) in enumerate(modules_data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = code
        row_cells[1].text = pages
        row_cells[2].text = desc

    p14 = doc.add_paragraph()
    p14.paragraph_format.space_before = Pt(12)
    p14.add_run("Nếu một module không nằm trong quyền truy cập của khách hàng, khi nhấp vào menu tương ứng, màn hình sẽ hiển thị cảnh báo từ chối truy cập màu đỏ và dừng toàn bộ tiến trình render để đảm bảo an toàn tuyệt đối.")
    p14.paragraph_format.space_after = Pt(12)

    # Lưu tệp tin Word vào thư mục gốc của dự án
    output_filename = "HUONG_DAN_SU_DUNG_VA_KICH_HOAT_LICENSE.docx"
    doc.save(output_filename)
    print(f"Build document success: {output_filename}")

if __name__ == "__main__":
    build_guide()
